import streamlit as st
import pandas as pd
import random
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from package.read_vocab import read_vocab


st.title('中考考纲词汇练习')

n_samples = st.slider('选择要抽取的词汇数量：', 5, 20, 10, 5)

df = read_vocab()
df_new = pd.DataFrame(columns=df.columns)

# 创建一个函数，该函数在点击按钮时触发文件下载
def download_file():
    with open("words.docx", "rb") as file:
        # 读取文件内容
        file_contents = file.read()
    # 显示下载链接
    st.download_button(
        "点击下载 words.docx",
        file_contents,
        file_name='words.docx',
        key="words-docx",
        help="点击按钮下载 words.docx 文件"
    )

if 'words_selected' not in st.session_state:
    st.session_state['words_selected'] = {}

if st.button('随机抽取默写词汇'):
    st.session_state['random_numbers'] = random.sample(range(len(df.index)), n_samples)

    for i in st.session_state['random_numbers']:
        df_new = pd.concat([df_new, df[i:i+1]])
    
    st.session_state['df_new'] = df_new

if 'df_new' in st.session_state:
    show_option = st.radio('选择是否显示中文翻译：', ['仅显示英文', '显示中英对照'], index=0, label_visibility='collapsed')
    if show_option == '仅显示英文':
        for index, row in st.session_state['df_new'].iterrows():
            st.markdown(f":arrow_right: {row['英文']}")
    else:
        for index, row in st.session_state['df_new'].iterrows():
            st.session_state['words_selected'][index] = st.checkbox(f":arrow_right: {row['英文']} -- {row['中文']}")

if 'df_new' in st.session_state and st.button('添加需要加强的词汇'):
    df_selected = pd.DataFrame(columns=df.columns)
    for key in st.session_state['words_selected'].keys():
        if st.session_state['words_selected'][key]:
            df_selected = pd.concat([df_selected, df[key:key+1]])

    df_selected = df_selected.reset_index()[['英文', '中文']]
    st.dataframe(df_selected, use_container_width=True)

if 'df_new' in st.session_state and st.button('创建生词表文档'):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(14)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)

    df_selected = pd.DataFrame(columns=df.columns)
    for key in st.session_state['words_selected'].keys():
        if st.session_state['words_selected'][key]:
            df_selected = pd.concat([df_selected, df[key:key+1]])

    for _, word in df_selected.iterrows():
        p = doc.add_paragraph()
        run = p.add_run(f"{word['英文']} ==> {word['中文']}")
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')

    doc.save('words.docx')
    download_file()
